"""Servicio para generar documentos de demanda Word a partir del modelo base."""
import os
import json
import re
import shutil
import tempfile
from datetime import datetime, timezone, timedelta
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from services.drive_service import upload_to_drive, get_drive_link
from services.sheets_tracking import load_tracking, save_tracking_entry
from docx.oxml.ns import qn

# Zona horaria Colombia (UTC-5)
COL_TZ = timezone(timedelta(hours=-5))

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(BASE_DIR, "static", "MODELO DE DEMANDA PARA AUTOMATIZAR.docx")
GENERATED_DIR = os.path.join(BASE_DIR, "generated_demandas")

_TEMP_TEMPLATE = os.path.join(tempfile.gettempdir(), "slh_demanda_modelo_v3.docx")

# Copy template to temp (handles OneDrive locking on Windows)
try:
    shutil.copy2(TEMPLATE_PATH, _TEMP_TEMPLATE)
except Exception:
    # On Render/Linux the file is accessible directly
    _TEMP_TEMPLATE = TEMPLATE_PATH

os.makedirs(GENERATED_DIR, exist_ok=True)

# Tabla de correos por conjunto para notificaciones de parte demandante
CORREOS_CONJUNTO = {
    "FLAMENCO": "flamencoparqueheredia@gmail.com",
    "BARLOVENTO": "conjuntobarlovento24@gmail.com",
    "MANATI": "Manatiedificio@gmail.com",
    "MANATÍ": "Manatiedificio@gmail.com",
    "CARACOLI": "conjuntoresidencialcaracoli@gmail.com",
    "CARACOLÍ": "conjuntoresidencialcaracoli@gmail.com",
    "MALIBU": "Malibuadmon1@gmail.com",
    "MALIBÚ": "Malibuadmon1@gmail.com",
    "CANDIL": "parqueherediacandil@gmail.com",
    "IGUAZU": "admonconjuntoiguazu2024@gmail.com",
    "IGUAZÚ": "admonconjuntoiguazu2024@gmail.com",
    "ORQUIDEA": "crjar.orquidea@gmail.com",
    "ORQUÍDEA": "crjar.orquidea@gmail.com",
}


def is_demanda_generated(cedula: str) -> bool:
    return cedula in load_tracking()


def get_demanda_info(cedula: str) -> dict | None:
    return load_tracking().get(cedula)


def get_all_generated() -> dict:
    return load_tracking()


def _get_correo_conjunto(conjunto: str) -> str:
    """Busca el correo del conjunto en la tabla de notificaciones."""
    key = conjunto.upper().strip()
    if key in CORREOS_CONJUNTO:
        return CORREOS_CONJUNTO[key]
    # Try partial match
    for k, v in CORREOS_CONJUNTO.items():
        if k in key or key in k:
            return v
    return ""


def _get_fecha_colombia() -> str:
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }
    now = datetime.now(COL_TZ)
    return f"{now.day} de {meses[now.month]} del {now.year}"


def _replace_in_paragraph_preserve_format(paragraph, replacements: dict):
    """
    Reemplaza placeholders en un párrafo preservando el formato.
    Maneja placeholders repartidos entre múltiples runs.
    """
    for old, new in replacements.items():
        full = "".join(r.text for r in paragraph.runs)
        if old not in full:
            continue

        # Build a map: for each character position, which run index owns it
        runs = paragraph.runs
        pos = 0
        run_boundaries = []  # list of (start_pos, end_pos, run_index)
        for i, run in enumerate(runs):
            length = len(run.text)
            run_boundaries.append((pos, pos + length, i))
            pos += length

        idx = full.find(old)
        while idx != -1:
            end_idx = idx + len(old)
            # Find which runs are affected
            affected = []
            for start, end, ri in run_boundaries:
                if start < end_idx and end > idx:
                    affected.append(ri)

            if affected:
                first_ri = affected[0]
                # Get text before placeholder in first run
                first_start = run_boundaries[first_ri][0]
                prefix = full[first_start:idx]
                # Get text after placeholder in last run
                last_ri = affected[-1]
                last_end = run_boundaries[last_ri][1]
                suffix = full[end_idx:last_end]
                # Set first affected run to prefix + new + suffix
                runs[first_ri].text = prefix + new + suffix
                # Clear other affected runs
                for ri in affected[1:]:
                    runs[ri].text = ""

            # Recalculate for next occurrence
            full = "".join(r.text for r in runs)
            idx = full.find(old)


def _replace_in_tables(doc, replacements: dict):
    """Reemplaza placeholders en todas las tablas."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_preserve_format(p, replacements)


def _set_obligation_tables(doc, hechos_data=None, pretensiones_data=None):
    """
    Limpia las tablas de obligaciones (Hechos=Table1, Pretensiones=Table2).
    """
    standard_concepts = [
        "Cuotas ordinarias de administración",
        "Intereses de mora",
        "Cuotas extraordinarias",
        "Gastos de cobranza",
    ]
    obligation_tables = []
    for table in doc.tables:
        if len(table.rows) > 0:
            header = table.rows[0].cells[0].text.strip()
            if "Concepto" in header:
                obligation_tables.append(table)

    for idx, table in enumerate(obligation_tables):
        data = None
        if idx == 0 and hechos_data:
            data = hechos_data
        elif idx == 1 and pretensiones_data:
            data = pretensiones_data

        # Remove existing data rows
        while len(table.rows) > 1:
            tr = table.rows[-1]._tr
            table._tbl.remove(tr)

        if data:
            for item in data:
                row = table.add_row()
                row.cells[0].text = item.get("concepto", "")
                row.cells[1].text = item.get("valor", "$0")
        else:
            for concept in standard_concepts:
                row = table.add_row()
                row.cells[0].text = concept
                row.cells[1].text = "$0"


def _remove_empty_paragraphs_between(doc, after_text: str, before_text: str, keep=1):
    """Elimina párrafos vacíos entre dos párrafos, dejando 'keep' vacíos."""
    paragraphs = doc.paragraphs
    after_idx = None
    before_idx = None
    for i, p in enumerate(paragraphs):
        if after_text in p.text and after_idx is None:
            after_idx = i
        if before_text in p.text and after_idx is not None:
            before_idx = i
            break
    if after_idx is not None and before_idx is not None:
        empty_indices = [
            i for i in range(after_idx + 1, before_idx)
            if not paragraphs[i].text.strip()
        ]
        # Remove all but 'keep' empty paragraphs (remove from end)
        to_remove = empty_indices[keep:]
        for i in reversed(to_remove):
            p_element = paragraphs[i]._p
            p_element.getparent().remove(p_element)


def _apply_keep_with_next(doc):
    """Aplica 'keep with next' a títulos y headings para evitar títulos huérfanos."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    heading_styles = {"Heading 1", "Heading 2", "Heading 3"}
    for p in doc.paragraphs:
        if p.style.name in heading_styles:
            ppr = p._p.find(_qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                p._p.insert(0, ppr)
            kwn = ppr.find(_qn("w:keepNext"))
            if kwn is None:
                kwn = OxmlElement("w:keepNext")
                ppr.append(kwn)


def _remove_all_highlights(doc):
    """Quita todos los highlights de todo el documento eliminando el XML directamente."""
    from docx.oxml.ns import qn as _qn

    def _strip_highlight(element):
        """Remove w:highlight from an rPr element."""
        rpr = element.find(_qn('w:rPr'))
        if rpr is not None:
            hl = rpr.find(_qn('w:highlight'))
            if hl is not None:
                rpr.remove(hl)

    def _process_paragraph(p):
        # Remove from paragraph-level rPr (pPr > rPr > highlight)
        ppr = p._p.find(_qn('w:pPr'))
        if ppr is not None:
            _strip_highlight(ppr)
        # Remove from each run
        for run in p.runs:
            run.font.highlight_color = None
            _strip_highlight(run._r)

    for p in doc.paragraphs:
        _process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p)


def generate_demanda(client: dict, hechos_data=None, pretensiones_data=None,
                     medida_cautelar_text: str = "") -> str:
    """
    Genera documento de demanda para un cliente.

    Args:
        client: dict con cedula, propietario, conjunto, torre, apto, mora, correo
        hechos_data: lista de dicts {concepto, valor} para tabla Hechos
        pretensiones_data: lista de dicts {concepto, valor} para tabla Pretensiones
        medida_cautelar_data: dict con campos de medida cautelar (fmi, ubicacion)

    Returns:
        Ruta al archivo generado.
    """
    doc = Document(_TEMP_TEMPLATE)

    propietario = client["propietario"]
    cedula = client["cedula"]
    row_id = client.get("row_id", cedula)
    conjunto = client["conjunto"]
    torre = client["torre"]
    apto = client["apto"]
    mora = client["mora"]
    correo = client.get("correo", "")
    fecha = _get_fecha_colombia()
    correo_conjunto = _get_correo_conjunto(conjunto)

    # Medida cautelar - free text
    medida_text = medida_cautelar_text.strip() if medida_cautelar_text else ""

    # --- STEP 1: Full paragraph text replacements for complex paragraphs ---
    # These paragraphs have placeholders split across many runs,
    # so we replace the full paragraph text.

    for i, p in enumerate(doc.paragraphs):
        full = p.text

        # Paragraph [56] - Notificaciones demandante
        if "Parte demandante:" in full and "{CONJUNTO" in full:
            if p.runs:
                # Clear all runs
                for r in p.runs:
                    r.text = ""
                # Run 0: "Parte demandante: " (bold=None, inherits heading bold)
                p.runs[0].text = "Parte demandante: "
                p.runs[0].bold = None
                # Run 1+: rest is not bold
                if len(p.runs) > 1:
                    p.runs[1].text = (
                        f"Dirección: Conjunto Residencial {conjunto}, "
                        f"ciudad de Cartagena, correos electrónicos: "
                        f"{correo_conjunto}"
                    )
                    p.runs[1].bold = False
            continue

        # Paragraph [60] - Notificaciones demandada
        if "Parte demandada:" in full and "Dirección" in full and "{CONJUNTO" in full:
            if p.runs:
                for r in p.runs:
                    r.text = ""
                # Run 0: "Parte demandada: " (bold=None, inherits heading bold)
                p.runs[0].text = "Parte demandada: "
                p.runs[0].bold = None
                # Run 1+: rest is not bold
                if len(p.runs) > 1:
                    if correo:
                        rest = (
                            f"Dirección: Conjunto Residencial {conjunto} "
                            f"Torre {torre} Apartamento {apto}, de la ciudad de "
                            f"Cartagena. Correo electrónico: {correo}"
                        )
                    else:
                        rest = (
                            f"Dirección: Conjunto Residencial {conjunto} "
                            f"Torre {torre} Apartamento {apto}, de la ciudad de "
                            f"Cartagena. Manifiesto bajo la gravedad de juramento "
                            f"que desconozco la dirección electrónica del demandado/ "
                            f"no tengo prueba que acredite propiedad del correo."
                        )
                    p.runs[1].text = rest
                    p.runs[1].bold = False
            continue

        # Paragraph [88] - Medida cautelar
        if "Embargo y secuestro" in full and "{060-17644}" in full:
            if medida_text and p.runs:
                p.runs[0].text = medida_text
                p.runs[0].bold = None
                p.runs[0].font.highlight_color = None
                for j in range(1, len(p.runs)):
                    p.runs[j].text = ""
                    p.runs[j].font.highlight_color = None
            continue

    # --- STEP 2: Run-level replacements for placeholders ---
    # These preserve bold/italic/font formatting per run
    replacements = {
        "{FECHA DE HOY}": fecha,
        "{CONJUNTO}": conjunto,
        "{CONJUNTO }": conjunto,
        "{PROPIETARIO}": propietario,
        "{CEDULA}": cedula,
        "{TORRE}": torre,
        "{APTO}": apto,
        "{CIUDAD}": "Cartagena",
        "27 de febrero del 2026": fecha,
    }

    for p in doc.paragraphs:
        _replace_in_paragraph_preserve_format(p, replacements)

    # Replace in tables (Table 0 = reference table)
    _replace_in_tables(doc, replacements)

    # --- STEP 3: Obligation tables ---
    _set_obligation_tables(doc, hechos_data, pretensiones_data)

    # --- STEP 3.5: Remove empty paragraphs between Segunda and Tercera ---
    _remove_empty_paragraphs_between(doc, "Segunda:", "Tercera:")

    # --- STEP 3.6: Keep headings with next paragraph (avoid orphan titles) ---
    _apply_keep_with_next(doc)

    # --- STEP 3.7: Adjust bottom margin to avoid text overlapping footer ---
    from docx.shared import Cm
    for section in doc.sections:
        section.bottom_margin = Cm(3.5)

    # --- STEP 4: Remove ALL highlights from the entire document ---
    _remove_all_highlights(doc)

    # --- STEP 5: Save ---
    safe_name = re.sub(r'[^\w\s-]', '', propietario).strip().replace(' ', '_')
    filename = f"Demanda_{safe_name}_{cedula}.docx"
    filepath = os.path.join(GENERATED_DIR, filename)
    doc.save(filepath)

    # Track in Google Sheets
    save_tracking_entry(row_id, {
        "propietario": propietario,
        "conjunto": conjunto,
        "filename": filename,
        "generated_at": datetime.now(COL_TZ).isoformat(),
        "mora": mora,
        "hechos": hechos_data or [],
        "pretensiones": pretensiones_data or [],
        "medida_cautelar": medida_text,
        "drive_id": "",
        "drive_link": "",
    })

    return filepath
